from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd

# Create a new Document
doc = Document()

# Header, make bold and centre
header = doc.add_heading(level=1)
header_run = header.add_run("A&E Attendances and Emergency Admissions\nJune 2024 Statistical Commentary")
header_run.bold = True
header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Line break
doc.add_paragraph()
findings_header = doc.add_heading(level=2)
findings_header_run = findings_header.add_run("Main Findings")
findings_header_run.bold = True

# Make new paragraph text
text = ("All growth rates are adjusted for calendar days. When comparing a month to the previous "
        "year, a daily average is used. Due to the Covid-19 response, caution should be exercised "
        "in drawing comparisons with data from previous years. "
        "From June 2023, 4 hour performance data from the 14 Clinical Review of Standard (CRS) "
        "field testing trusts has been reintroduced. Care should be taken when comparing "
        "performance during the field-testing period (May 2019 – June 2023). For further "
        "information on the impact of reintroducing the field-testing trust please see the note here: "
        "https://www.england.nhs.uk/statistics/statistical-work-areas/ae-waiting-times-and-activity/")
paragraph = doc.add_paragraph(text)

doc.add_paragraph()
attendances_header = doc.add_heading(level=2)
attendances_header_run = attendances_header.add_run("Attendances")
attendances_header_run.bold = True


# Load csv data
# Load in, filter data, and do calculations
data = pd.read_csv('total_attendances.csv')
england_june_2023 = data[(data['month_year'] == 'June 2023') & (data['region'] == 'England')]
england_june_2024 = data[(data['month_year'] == 'June 2024') & (data['region'] == 'England')]
percentage_change = (((england_june_2024['total'].values[0] - england_june_2023['total'].values[0]) / england_june_2024['total'].values[0]) * 100)


# Create the summary sentence
summary_sentence = (
    f"The total number of attendances in June 2024 was {england_june_2024['total'].values[0]}. "
    f"This is an increase of {percentage_change:.2f}% when comparing the daily average attendances to June 2023."
)
doc.add_paragraph(summary_sentence, style='List Bullet')






# Save the document
doc.save('statistical_commentary.docx')
